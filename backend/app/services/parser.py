import fitz
import docx
import pandas as pd
import pytesseract

from PIL import Image
from pathlib import Path
from typing import List, Dict


class DocumentParser:

    def parse(self, filepath: Path) -> List[Dict]:

        ext = filepath.suffix.lower()

        print(f"\nParsing file: {filepath.name}")
        print(f"Extension: {ext}")

        if ext == ".pdf":
            pages = self._parse_pdf(filepath)

        elif ext == ".docx":
            pages = self._parse_docx(filepath)

        elif ext == ".txt":
            pages = self._parse_txt(filepath)

        elif ext == ".csv":
            pages = self._parse_csv(filepath)

        elif ext == ".xlsx":
            pages = self._parse_xlsx(filepath)

        elif ext in [".png", ".jpg", ".jpeg"]:
            pages = self._parse_image(filepath)

        else:
            raise ValueError(f"Unsupported file type: {ext}")

        print(f"Pages extracted: {len(pages)}")

        for i, page in enumerate(pages):
            print(
                f"Page {i+1}: {len(page.get('text',''))} characters"
            )

        return pages

    # -----------------------------------------------------

    def _parse_pdf(self, filepath: Path) -> List[Dict]:

        pages = []

        doc = fitz.open(filepath)

        for i, page in enumerate(doc):

            text = page.get_text("text").strip()

            if len(text) < 20:

                pix = page.get_pixmap(dpi=200)

                img = Image.frombytes(
                    "RGB",
                    [pix.width, pix.height],
                    pix.samples,
                )

                text = pytesseract.image_to_string(img).strip()

            if text:

                pages.append({
                    "page_number": i + 1,
                    "text": text
                })

        doc.close()

        return pages

    # -----------------------------------------------------

    def _parse_docx(self, filepath: Path) -> List[Dict]:

        document = docx.Document(filepath)

        text = []

        for para in document.paragraphs:

            if para.text.strip():
                text.append(para.text.strip())

        for table in document.tables:

            for row in table.rows:

                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )

                if row_text:
                    text.append(row_text)

        final_text = "\n".join(text).strip()

        print(f"DOCX extracted {len(final_text)} characters")

        if not final_text:
            return []

        return [
            {
                "page_number": 1,
                "text": final_text,
            }
        ]

    # -----------------------------------------------------

    def _parse_txt(self, filepath: Path):

        text = filepath.read_text(
            encoding="utf-8",
            errors="ignore",
        ).strip()

        if not text:
            return []

        return [{
            "page_number": 1,
            "text": text,
        }]

    # -----------------------------------------------------

    def _parse_csv(self, filepath: Path):

        df = pd.read_csv(filepath)

        text = df.to_string(index=False)

        if not text.strip():
            return []

        return [{
            "page_number": 1,
            "text": text,
        }]

    # -----------------------------------------------------

    def _parse_xlsx(self, filepath: Path):

        excel = pd.ExcelFile(filepath)

        pages = []

        for sheet in excel.sheet_names:

            df = excel.parse(sheet)

            text = df.to_string(index=False)

            if text.strip():

                pages.append({
                    "page_number": sheet,
                    "text": text,
                })

        return pages

    # -----------------------------------------------------

    def _parse_image(self, filepath: Path):

        text = pytesseract.image_to_string(
            Image.open(filepath)
        ).strip()

        if not text:
            return []

        return [{
            "page_number": 1,
            "text": text,
        }]