from google import genai
from sqlalchemy.orm import Session
from datetime import datetime
import json

from app.config import GEMINI_API_KEY, LLM_MODEL, DATA_DIR
from app.services.maintenance import MaintenanceIntelligence
from app.services.compliance import ComplianceEngine
from app.services.rag import RAGEngine

client = genai.Client(api_key=GEMINI_API_KEY)

REPORTS_DIR = DATA_DIR / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

REPORT_PROMPTS = {
    "maintenance": """Write a formal Maintenance Report section-by-section using the provided data.
Include: Overview, Equipment Health Summary, Key Failures, Downtime Analysis, Recommendations.
Be specific with numbers from the data. Do not invent figures not present in the data.""",
    "incident": """Write a formal Incident Report based on the provided incident/RCA data.
Include: Incident Summary, Root Cause, Corrective Actions Taken, Preventive Recommendations, Sign-off section.""",
    "audit_summary": """Write a formal Audit Summary Report using the provided compliance data.
Include: Overview, Compliance Rate, Non-Conformances (missing/expired), Risk Assessment, Action Items.""",
    "equipment_summary": """Write a formal Equipment Summary Report for the specified equipment using the provided data.
Include: Equipment Overview, Maintenance History, Failure Trends, Current Compliance Status, Recommendations.""",
    "executive": """Write a concise Executive Report (1-2 pages) summarizing plant-wide status using the provided data.
Include: Executive Summary, Key Metrics, Top Risks, Compliance Snapshot, Strategic Recommendations.
Keep language high-level and suited for leadership, not floor engineers.""",
}


def _extract_json_from_response(text: str) -> dict:
    """Extract JSON from Gemini response, stripping markdown code fences if present."""
    if not text:
        return {}
    cleaned = text.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split("\n", 1)
        cleaned = lines[1] if len(lines) > 1 else cleaned
        if "```" in cleaned:
            cleaned = cleaned.rsplit("```", 1)[0]
    cleaned = cleaned.strip()
    return json.loads(cleaned)


class ReportGenerator:
    def __init__(self):
        self.maintenance = MaintenanceIntelligence()
        self.compliance = ComplianceEngine()
        self.rag = RAGEngine()

    def _gather_data(self, db: Session, report_type: str, equipment_tag: str = None) -> dict:
        data = {}
        if report_type in ("maintenance", "executive", "equipment_summary"):
            data["health_overview"] = self.maintenance.equipment_health_overview(db)
            data["predictive_alerts"] = self.maintenance.predictive_alerts(db)
        if report_type in ("audit_summary", "executive", "equipment_summary"):
            data["compliance_summary"] = self.compliance.dashboard_summary(db)
        if report_type == "equipment_summary" and equipment_tag:
            data["failure_trends"] = self.maintenance.failure_trends(db, equipment_tag)
            data["equipment_health"] = [
                e for e in data["health_overview"] if e["equipment_tag"] == equipment_tag
            ]
        return data

    def generate(self, db: Session, report_type: str, equipment_tag: str = None,
                 incident_description: str = None) -> dict:
        if report_type not in REPORT_PROMPTS:
            raise ValueError(f"Unknown report type: {report_type}")

        if report_type == "incident":
            if not incident_description:
                raise ValueError("incident_description is required for incident reports")
            from app.services.root_cause import RootCauseAnalyzer
            rca_result = RootCauseAnalyzer().analyze(incident_description)
            data = {"incident_description": incident_description, "rca": rca_result}
        else:
            data = self._gather_data(db, report_type, equipment_tag)

        system_prompt = REPORT_PROMPTS[report_type] + \
            "\n\nReturn STRICT JSON: {\"title\": str, \"sections\": [{\"heading\": str, \"content\": str}]}"
        user_prompt = f"Data:\n{json.dumps(data, default=str, indent=2)}"
        full_prompt = f"{system_prompt}\n\n{user_prompt}"

        response = client.models.generate_content(
            model=LLM_MODEL,
            contents=full_prompt,
        )

        text = response.text if hasattr(response, "text") and response.text else "{}"
        report_content = _extract_json_from_response(text)
        report_content["report_type"] = report_type
        report_content["generated_at"] = datetime.utcnow().isoformat()
        return report_content

    def export_pdf(self, report: dict):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch

        filename = f"{report['report_type']}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.pdf"
        filepath = REPORTS_DIR / filename

        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("TitleStyle", parent=styles["Title"], spaceAfter=20)
        heading_style = ParagraphStyle("Heading", parent=styles["Heading2"], spaceBefore=14, spaceAfter=8)

        story = [Paragraph(report["title"], title_style),
                 Paragraph(f"Generated: {report['generated_at']}", styles["Normal"]),
                 Spacer(1, 0.2 * inch)]

        for section in report["sections"]:
            story.append(Paragraph(section["heading"], heading_style))
            for para in section["content"].split("\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), styles["Normal"]))
                    story.append(Spacer(1, 0.1 * inch))

        doc.build(story)
        return filepath

    def export_docx(self, report: dict):
        from docx import Document as DocxDocument
        from docx.shared import Pt

        filename = f"{report['report_type']}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
        filepath = REPORTS_DIR / filename

        doc = DocxDocument()
        doc.add_heading(report["title"], level=0)
        meta = doc.add_paragraph(f"Generated: {report['generated_at']}")
        meta.runs[0].italic = True

        for section in report["sections"]:
            doc.add_heading(section["heading"], level=1)
            for para in section["content"].split("\n"):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.style.font.size = Pt(11)

        doc.save(filepath)
        return filepath

